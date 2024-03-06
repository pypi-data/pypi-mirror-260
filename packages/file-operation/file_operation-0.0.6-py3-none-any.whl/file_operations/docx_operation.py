# -*- coding: utf-8 -*-
"""
@Time : 2023/8/16 21:09 
@Author : skyoceanchen
@TEL: 18916403796
@项目：文件使用
@File : docx_operation.by
@PRODUCT_NAME :PyCharm
"""
import docx  # pip install python-docx
from docx import Document  # 用于新建文档
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 对齐库
from docx.enum.table import WD_TABLE_ALIGNMENT  # 对齐库
from docx.shared import Pt  # 导入字号，颜色库，注意大小写
from docx.shared import RGBColor  # 颜色库，注意大小写
from docx.oxml.ns import qn  # 导入中文格式库
from docx.enum.text import WD_LINE_SPACING  # docx文档的行间距模式设置常量
from docx.shared import Cm  # docx文档页边距
import time  # 用来确定时间
import datetime
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE


class DocxOperation(object):
    def __init__(self, path):
        self.path = path

    def read_docx(self):
        doc = docx.Document(self.path)  # 绝对路径
        # 读取表格外全部内容
        text_list = []
        for i in doc.paragraphs:
            text = i.text.replace("—", "")
            if text:
                text_list.append(text)
        text = ''.join(text_list)
        return text


class DocxOperationCreate(object):
    def __init__(self):
        self.document = Document()

    def global_config(self):
        # 向document文档写入数据的过程
        # 设置全局样式
        self.document.styles['Normal'].font.name = 'Times New Roman'  # 英文和数字的字体设置；也可以直接设置仿宋_GB2312等中文格式
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),
                                                               '仿宋_GB2312')  # 指定全文非标题（即正文）中文字体进行格式设置。仿宋_GB2312
        self.document.styles['Normal'].font.size = Pt(16)  # 字号    14：四号；16：三号；22号：二号
        self.document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)  # 颜色黑色
        self.document.styles['Normal'].paragraph_format.line_spacing = Pt(30)  # 若=1则行距为1行；若=docx.shared.Pt(30)则行距为30磅
        self.document.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后距离0磅
        self.document.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前距离0磅
        self.document.sections[0].top_margin = Cm(3.0)  # 新建一个空白文档，那么默认是只有一个节，如果这里没有[0],会报错
        self.document.sections[0].bottom_margin = Cm(3.0)
        self.document.sections[0].left_margin = Cm(2.6)
        self.document.sections[0].right_margin = Cm(2.6)
        # print('上',self.document.sections[0].top_margin.cm,'下',self.document.sections[0].bottom_margin.cm,'左',self.document.sections[0].left_margin.cm,'右',self.document.sections[0].right_margin.cm)
        # 也可以用下面这个表达式，注意如果后面不加.cm，输出的数字就不是cm
        print('上%.2fcm,下%.2fcm,左%.2fcm,右%.2fcm.' % (
            self.document.sections[0].top_margin, self.document.sections[0].bottom_margin.cm,
            self.document.sections[0].left_margin.cm,
            self.document.sections[0].right_margin.cm))

    def add_heading(self, text, level=0, alignment=1):
        """
        :param title:
        :param level: level 标题等级， 0-9
        :param alignment: 1 为居中 0 左 2右 0，1，2，3，4，5，7，8，9
        :return:
        """
        heading = self.document.add_heading(text, level)
        heading.alignment = alignment
        print(heading, dir(heading))
        print(heading.style, dir(heading.style))
        # self.font(heading.style, bold=True, italic=True, underline=True, strike=True, shadow=True, size=30)
        # heading.style.font.size = Pt(100)
        heading.style.font.name = "微软雅黑"
        print(heading._element)
        heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 中文字体
        # heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")  # 中文字体

    # _element
    # r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    def font_run(self,
                 obj, bold=False, italic=False, underline=False, strike=False, shadow=False, size=None,
                 rgb=(0, 0, 0), font_name=None, alignment=None
                 ):
        """

        :param obj:  paragraph.add_run(title='HWD 标定报告', )
        :param bold:
        :param italic:
        :param underline:
        :param strike:
        :param shadow:
        :param size:
        :param rgb:
        :param font_name:
        :param alignment:
        :return:
        """
        obj.font.bold = bold  # 加粗
        obj.font.italic = italic  # 斜体
        obj.font.underline = underline  # 下划线
        obj.font.strike = strike  # 删除线
        obj.font.shadow = shadow  # 阴影
        obj.font.color.rgb = RGBColor(*rgb)
        if size:
            obj.font.size = Pt(size)  # 字体大小
        if font_name:
            obj.font.name = font_name
            obj._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
        if alignment or alignment == 0:
            obj._parent.alignment = alignment
        return obj

    def font_paragraph(self, paragraph, bold=False, italic=False, underline=False, strike=False, shadow=False,
                       size=None, rgb=(0, 0, 0), font_name=None, alignment=None,
                       line_spacing=0, space_after=0, space_before=0, first_line_indent=0
                       ):
        obj = paragraph.style
        obj.font.bold = bold  # 加粗
        obj.font.italic = italic  # 斜体
        obj.font.underline = underline  # 下划线
        obj.font.strike = strike  # 删除线
        obj.font.shadow = shadow  # 阴影
        obj.font.color.rgb = RGBColor(*rgb)
        if size:
            obj.font.size = Pt(size)  # 字体大小
        if obj.font.size:
            size = obj.font.size
        else:
            size = 11
        if font_name:
            obj.font.name = font_name
            obj._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
        if alignment or alignment == 0:
            paragraph.alignment = alignment
        if line_spacing or line_spacing == 0:
            paragraph.paragraph_format.line_spacing = Pt(line_spacing)  # 若=1则行距为1行；若=docx.shared.Pt(30)则行距为30磅
        if space_after or space_after == 0:
            paragraph.paragraph_format.space_after = Pt(space_after)  # 段后距离0磅
        if space_before or space_before == 0:
            paragraph.paragraph_format.space_before = Pt(space_before)  # 段前距离0磅
        if first_line_indent:
            paragraph.paragraph_format.first_line_indent = Pt(size * first_line_indent)  # 2 # 设置段落的缩进量
        return paragraph

    def add_paragraph(self, conent=""):
        paragraph = self.document.add_paragraph(conent)
        return paragraph

    def add_run(self, paragraph=None, conent=""):
        if paragraph:
            run = paragraph.add_run(conent)
        else:
            run = self.document.add_paragraph().add_run(conent)
        return run

    def add_table(self, rows, cols):
        table = self.document.add_table(rows=rows, cols=cols)
        return table

    def table_merge(self, table, start, end):
        """
        :param table:
        :param start: 开始坐标
        :param end: 结束坐标
        :return:
        table.cell(0, 0).merge(table.cell(0, 1))
        """
        table.cell(*start).merge(table.cell(*end))
        return table
        #

    def table_content(self, table, content):
        """

        :param table:
        :param content:
        :return:
        """
        # 循环添加
        for i, row in enumerate(table.rows):

            for j, cell in enumerate(row.cells):
                cell.text = content[i][j]
                # 获取单元格中的段落对象
                # paragraph = cell.paragraphs[0]
                # # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # # 这里的run可以设置一些属性
                # run = paragraph.add_run(content[i][j])
                # # if i == 2:
                # #     run.font.color.rgb = RGBColor(0, 255, 0)

    def font_table(self, table,
                   bold=False, italic=False, underline=False, strike=False, shadow=False, size=None,
                   rgb=(0, 0, 0), font_name=None, alignment=None
                   ):
        """

        :param table:  document.add_table(rows=rows, cols=cols)
        :param bold:
        :param italic:
        :param underline:
        :param strike:
        :param shadow:
        :param size:
        :param rgb:
        :param font_name:
        :param alignment:
        :return:
        """
        obj = table.style
        obj.font.bold = bold  # 加粗
        obj.font.italic = italic  # 斜体
        obj.font.underline = underline  # 下划线
        obj.font.strike = strike  # 删除线
        obj.font.shadow = shadow  # 阴影
        obj.font.color.rgb = RGBColor(*rgb)
        if size:
            obj.font.size = Pt(size)  # 字体大小
        if font_name:
            obj.font.name = font_name
            obj._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
        if alignment:
            table.style.paragraph_format.alignment = alignment

    def table_div_style(self, table_style_name='mystyle', bold=False, italic=False, underline=False, strike=False,
                        shadow=False, size=None, rgb=(0, 0, 0), font_name=None, alignment=None):
        """

        :param table_style_name:
        :param bold:
        :param font_name:
        :param size:
        :param rgb:
        :param alignment:
        :return:

        用法：
         # 创建一个带有三列和四行的表格
        # table = document.add_table(rows=4, cols=3,style='mystyle')
        table = document.add_table(rows=4, cols=3)
        # 设置表格样式
        table.style = obj
        """
        # 定义样式
        table_style_name = table_style_name
        obj = self.document.styles.add_style(table_style_name, WD_STYLE_TYPE.TABLE)
        obj.font.bold = bold  # 加粗
        obj.font.italic = italic  # 斜体
        obj.font.underline = underline  # 下划线
        obj.font.strike = strike  # 删除线
        obj.font.shadow = shadow  # 阴影
        obj.font.color.rgb = RGBColor(*rgb)
        if alignment or alignment == 0:
            obj.paragraph_format.alignment = alignment
        if size:
            obj.font.size = Pt(size)  # 字体大小
        if font_name:
            obj.font.name = font_name
            obj._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
        return obj

    def table_global_config(self, table):

        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 对表格进行对齐
        # Turn off autofit  固定列宽
        table.autofit = False

    # 换行
    def add_break(self):
        self.font_paragraph(self.document.add_paragraph("\t"))

    # 分页
    def add_page_break(self):
        self.document.add_page_break()
